"""
Export Service for AI SEO Agent System
Handles content export in multiple formats (PDF, DOCX, MD, HTML, JSON)
"""

import io
import json
import logging
from typing import Any, Dict, Optional
from pathlib import Path
from uuid import UUID

from sqlalchemy.ext.asyncio import AsyncSession

from backend.db.models import ContentVersion
from backend.services.storage import StorageService
from backend.core.logging import get_logger

logger = get_logger(__name__)


class ExportService:
    """Service for exporting content in various formats"""

    def __init__(self):
        self.storage = StorageService()

    async def export_content(
        self,
        session: AsyncSession,
        content_id: UUID,
        format: str = "json"
    ) -> Dict[str, Any]:
        """
        Export content in specified format
        
        Args:
            session: Database session
            content_id: Content version UUID
            format: Export format (pdf, docx, md, html, json)
            
        Returns:
            Export result with file URL
        """
        # Get content version
        from sqlalchemy import select
        stmt = select(ContentVersion).where(ContentVersion.id == content_id)
        result = await session.execute(stmt)
        content = result.scalar_one_or_none()
        
        if not content:
            return {"success": False, "error": "Content not found"}
        
        # Fetch content from storage
        content_data = await self._fetch_content(content)
        
        # Generate export based on format
        format_handlers = {
            "json": self._export_json,
            "md": self._export_markdown,
            "html": self._export_html,
            "docx": self._export_docx,
            "pdf": self._export_pdf,
        }
        
        handler = format_handlers.get(format.lower())
        if not handler:
            return {"success": False, "error": f"Unsupported format: {format}"}
        
        try:
            file_buffer, filename = await handler(content, content_data)
            
            # Upload to storage
            upload_key = f"exports/{content_id}/{filename}"
            file_url = await self.storage.upload_file(
                bucket="exports",
                key=upload_key,
                file_buffer=file_buffer,
                content_type=self._get_content_type(format)
            )
            
            logger.info(f"Content exported successfully: {filename}")
            
            return {
                "success": True,
                "filename": filename,
                "url": file_url,
                "format": format,
                "content_id": str(content_id)
            }
            
        except Exception as e:
            logger.error(f"Export failed: {str(e)}")
            return {"success": False, "error": str(e)}

    async def _fetch_content(self, content: ContentVersion) -> str:
        """Fetch content from storage"""
        if content.s3_key:
            return await self.storage.download_file(
                bucket="content",
                key=content.s3_key
            )
        return ""

    async def _export_json(
        self, 
        content: ContentVersion, 
        content_data: str
    ) -> tuple[io.BytesIO, str]:
        """Export as JSON"""
        data = {
            "id": str(content.id),
            "title": content.title,
            "meta_description": content.meta_description,
            "slug": content.slug,
            "target_keyword": content.target_keyword,
            "secondary_keywords": content.secondary_keywords,
            "content": content_data,
            "word_count": content.word_count,
            "semantic_score": content.semantic_score,
            "readability_score": content.readability_score,
            "created_at": content.created_at.isoformat(),
            "status": content.status.value
        }
        
        buffer = io.BytesIO()
        buffer.write(json.dumps(data, indent=2).encode('utf-8'))
        buffer.seek(0)
        
        return buffer, f"{content.slug or 'content'}.json"

    async def _export_markdown(
        self, 
        content: ContentVersion, 
        content_data: str
    ) -> tuple[io.BytesIO, str]:
        """Export as Markdown"""
        # If content is HTML, convert to markdown
        if content_data.strip().startswith('<'):
            try:
                import markdown
                # Simple HTML to MD conversion (basic)
                md_content = self._html_to_markdown(content_data)
            except ImportError:
                md_content = content_data
        else:
            md_content = content_data
        
        # Add frontmatter
        frontmatter = f"""---
title: "{content.title or ''}"
meta_description: "{content.meta_description or ''}"
slug: "{content.slug or ''}"
target_keyword: "{content.target_keyword or ''}"
---

"""
        full_content = frontmatter + md_content
        
        buffer = io.BytesIO()
        buffer.write(full_content.encode('utf-8'))
        buffer.seek(0)
        
        return buffer, f"{content.slug or 'content'}.md"

    async def _export_html(
        self, 
        content: ContentVersion, 
        content_data: str
    ) -> tuple[io.BytesIO, str]:
        """Export as HTML"""
        # If already HTML, use as-is
        if content_data.strip().startswith('<'):
            html_content = content_data
        else:
            # Convert markdown to HTML
            try:
                import markdown
                html_content = markdown.markdown(
                    content_data,
                    extensions=['extra', 'codehilite', 'toc']
                )
                # Wrap in basic HTML template
                html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{content.title or 'Content'}</title>
    <meta name="description" content="{content.meta_description or ''}">
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, Oxygen, Ubuntu, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }}
        h1, h2, h3 {{ color: #1a1a1a; }}
        a {{ color: #0066cc; }}
        code {{ background: #f4f4f4; padding: 2px 6px; border-radius: 3px; }}
        pre {{ background: #f4f4f4; padding: 15px; border-radius: 5px; overflow-x: auto; }}
    </style>
</head>
<body>
    <article>
        {html_content}
    </article>
</body>
</html>"""
            except ImportError:
                html_content = f"<p>{content_data}</p>"
        
        buffer = io.BytesIO()
        buffer.write(html_content.encode('utf-8'))
        buffer.seek(0)
        
        return buffer, f"{content.slug or 'content'}.html"

    async def _export_docx(
        self, 
        content: ContentVersion, 
        content_data: str
    ) -> tuple[io.BytesIO, str]:
        """Export as DOCX"""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title
            if content.title:
                doc.add_heading(content.title, level=1)
            
            # Add meta description
            if content.meta_description:
                doc.add_paragraph(content.meta_description, style='Intense Quote')
            
            # Add content
            if content_data.strip().startswith('<'):
                # HTML to DOCX (basic conversion)
                paragraphs = self._html_to_paragraphs(content_data)
                for para in paragraphs:
                    doc.add_paragraph(para)
            else:
                # Plain text or markdown
                for line in content_data.split('\n\n'):
                    if line.strip():
                        doc.add_paragraph(line)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer, f"{content.slug or 'content'}.docx"
            
        except ImportError:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")

    async def _export_pdf(
        self, 
        content: ContentVersion, 
        content_data: str
    ) -> tuple[io.BytesIO, str]:
        """Export as PDF"""
        try:
            from weasyprint import HTML, CSS
            
            # Generate HTML first
            html_buffer, _ = await self._export_html(content, content_data)
            html_content = html_buffer.read().decode('utf-8')
            
            # Add print-specific CSS
            css = CSS(string="""
                @page {
                    size: A4;
                    margin: 2cm;
                }
                body {
                    font-family: Georgia, serif;
                    line-height: 1.6;
                    color: #333;
                }
                h1, h2, h3 {
                    color: #1a1a1a;
                    page-break-after: avoid;
                }
                p {
                    text-align: justify;
                }
                a {
                    color: #0066cc;
                    text-decoration: none;
                }
                @media print {
                    a[href]::after {
                        content: " (" attr(href) ")";
                    }
                }
            """)
            
            # Generate PDF
            pdf_doc = HTML(string=html_content)
            buffer = io.BytesIO()
            pdf_doc.write_pdf(buffer, stylesheets=[css])
            buffer.seek(0)
            
            return buffer, f"{content.slug or 'content'}.pdf"
            
        except ImportError:
            raise ImportError("weasyprint not installed. Install with: pip install weasyprint")

    def _html_to_markdown(self, html: str) -> str:
        """Basic HTML to Markdown conversion"""
        import re
        
        md = html
        
        # Headers
        md = re.sub(r'<h1[^>]*>(.*?)</h1>', r'# \1', md, flags=re.DOTALL)
        md = re.sub(r'<h2[^>]*>(.*?)</h2>', r'## \1', md, flags=re.DOTALL)
        md = re.sub(r'<h3[^>]*>(.*?)</h3>', r'### \1', md, flags=re.DOTALL)
        
        # Bold and italic
        md = re.sub(r'<strong[^>]*>(.*?)</strong>', r'**\1**', md, flags=re.DOTALL)
        md = re.sub(r'<b[^>]*>(.*?)</b>', r'**\1**', md, flags=re.DOTALL)
        md = re.sub(r'<em[^>]*>(.*?)</em>', r'*\1*', md, flags=re.DOTALL)
        md = re.sub(r'<i[^>]*>(.*?)</i>', r'*\1*', md, flags=re.DOTALL)
        
        # Links
        md = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', md, flags=re.DOTALL)
        
        # Lists
        md = re.sub(r'<li[^>]*>(.*?)</li>', r'- \1', md, flags=re.DOTALL)
        
        # Code
        md = re.sub(r'<code[^>]*>(.*?)</code>', r'`\1`', md, flags=re.DOTALL)
        md = re.sub(r'<pre[^>]*>(.*?)</pre>', r'```\n\1\n```', md, flags=re.DOTALL)
        
        # Remove remaining tags
        md = re.sub(r'<[^>]+>', '', md)
        
        # Clean up whitespace
        md = re.sub(r'\n\s*\n', '\n\n', md)
        
        return md.strip()

    def _html_to_paragraphs(self, html: str) -> list[str]:
        """Extract paragraphs from HTML"""
        import re
        
        # Extract text from common block elements
        paragraphs = []
        
        # Match p, div, li tags
        pattern = r'<(?:p|div|li)[^>]*>(.*?)</(?:p|div|li)>'
        matches = re.findall(pattern, html, re.DOTALL)
        
        for match in matches:
            # Remove inner tags
            text = re.sub(r'<[^>]+>', '', match)
            text = text.strip()
            if text:
                paragraphs.append(text)
        
        return paragraphs if paragraphs else [html]

    def _get_content_type(self, format: str) -> str:
        """Get MIME type for format"""
        types = {
            "json": "application/json",
            "md": "text/markdown",
            "html": "text/html",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pdf": "application/pdf",
        }
        return types.get(format.lower(), "application/octet-stream")
